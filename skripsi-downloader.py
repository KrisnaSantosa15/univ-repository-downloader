import os
import requests
from PIL import Image
from io import BytesIO
import img2pdf
import pytesseract
from docx import Document
import time

# ========== KONFIGURASI ==========
COOKIE_STRING = "your_cookie_here"
# jangan panggil/print cookie ini di log jika berbagi script

# Range halaman (sesuaikan jika perlu)
BAB2_START, BAB2_END = 0, 20
BAB4_START, BAB4_END = 0, 20

# Base URL (saya gunakan pola img; script akan coba 'file' jika 404)
BASE_BAB2 = "https://reader-repository.upi.edu/index.php/display/img/130155/3/"
BASE_BAB4_IMG = "https://reader-repository.upi.edu/index.php/display/img/130155/4/"
BASE_BAB4_FILE = "https://reader-repository.upi.edu/index.php/display/file/130155/4/"

OUTDIR = "downloaded"

# ================================

headers = {
    "Cookie": COOKIE_STRING,
    "User-Agent": "Mozilla/5.0 (compatible; script/1.0)"
}

os.makedirs(OUTDIR, exist_ok=True)


def download_range(base_url_img, base_url_file, start, end, outsub):
    path = os.path.join(OUTDIR, outsub)
    os.makedirs(path, exist_ok=True)
    saved_files = []
    for i in range(start, end + 1):
        # coba pola img dulu, jika 404 coba pola file (ada variasi di link yang kamu sebut)
        tried_urls = [f"{base_url_img}{i}", f"{base_url_file}{i}"]
        content = None
        used_url = None
        for url in tried_urls:
            try:
                r = requests.get(url, headers=headers,
                                 timeout=20, verify=False)
            except Exception as e:
                print(f"Error request {url}: {e}")
                continue
            if r.status_code == 200 and r.content:
                content = r.content
                used_url = url
                break
            else:
                # beberapa server merespon 200 tapi content html (login page). Cek content-type
                ct = r.headers.get("Content-Type", "")
                if r.status_code == 200 and ("image" in ct or r.content[:4].startswith(b'\xff\xd8') or r.content[:8].startswith(b'\x89PNG')):
                    content = r.content
                    used_url = url
                    break
                # else try next pattern
        if content is None:
            print(f"[!] Gagal mengunduh halaman {i} (dicoba: {tried_urls}).")
            continue
        # Simpan file dengan ekstensi berdasarkan header (fallback .jpg)
        fname = f"page_{i:03d}.jpg"
        # coba deteksi png
        if content[:8].startswith(b'\x89PNG'):
            fname = f"page_{i:03d}.png"
        elif content[:2] == b'BM':
            fname = f"page_{i:03d}.bmp"
        with open(os.path.join(path, fname), "wb") as f:
            f.write(content)
        saved_files.append(os.path.join(path, fname))
        print(f"Unduh: {used_url} -> {outsub}/{fname}")
        time.sleep(0.2)  # jeda kecil agar tidak membanjiri server
    return saved_files


def images_to_pdf(img_files, out_pdf):
    # pastikan urut
    img_files = sorted(img_files)
    with open(out_pdf, "wb") as f:
        f.write(img2pdf.convert(img_files))
    print(f"PDF tersimpan: {out_pdf}")


def ocr_images_to_docx(img_files, out_docx, lang="ind"):
    doc = Document()
    for img_path in sorted(img_files):
        try:
            img = Image.open(img_path)
        except Exception as e:
            print(f"Cannot open {img_path}: {e}")
            continue
        # OCR - atur bahasa sesuai instalasi tesseract (mis. 'ind' untuk Bahasa Indonesia)
        try:
            text = pytesseract.image_to_string(img, lang=lang)
        except Exception as e:
            print(f"OCR failed for {img_path}: {e}")
            text = ""
        doc.add_heading(os.path.basename(img_path), level=3)
        for line in text.splitlines():
            doc.add_paragraph(line)
    doc.save(out_docx)
    print(f"OCR DOCX tersimpan: {out_docx}")


if __name__ == "__main__":
    # BAB 2
    bab2_imgs = download_range(
        BASE_BAB2, BASE_BAB2, BAB2_START, BAB2_END, "bab2")
    if bab2_imgs:
        images_to_pdf(bab2_imgs, os.path.join(OUTDIR, "BAB2.pdf"))
        # ocr_images_to_docx(bab2_imgs, os.path.join(
        #     OUTDIR, "BAB2.docx"), lang="ind")
    # BAB 4
    bab4_imgs = download_range(
        BASE_BAB4_IMG, BASE_BAB4_FILE, BAB4_START, BAB4_END, "bab4")
    if bab4_imgs:
        images_to_pdf(bab4_imgs, os.path.join(OUTDIR, "BAB4.pdf"))
        # ocr_images_to_docx(bab4_imgs, os.path.join(
        #     OUTDIR, "BAB4.docx"), lang="ind")

    print("Selesai.")
