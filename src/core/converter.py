"""
Converter Module for Skripsi Downloader
Handles conversion of images to PDF and DOCX
"""
import os
from PIL import Image
import img2pdf
import pytesseract
from docx import Document
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn

console = Console()


class FileConverter:
    """Class to handle file conversions"""

    def __init__(self, config):
        self.config = config

    def images_to_pdf(self, image_files, output_path):
        """
        Convert a list of image files to PDF

        Args:
            image_files: List of image file paths
            output_path: Output PDF file path

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if not image_files:
                console.print("[red]No image files to convert![/red]")
                return False

            console.print(
                f"\n[bold cyan]Converting {len(image_files)} images to PDF...[/bold cyan]")

            # Sort files to ensure correct order
            sorted_files = sorted(image_files)

            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                console=console
            ) as progress:
                task = progress.add_task("[cyan]Processing images...", total=1)

                with open(output_path, "wb") as f:
                    f.write(img2pdf.convert(sorted_files))

                progress.update(task, advance=1)

            console.print(
                f"[bold green]✓ PDF saved:[/bold green] {output_path}\n")
            return True

        except Exception as e:
            console.print(f"[bold red]✗ Error creating PDF:[/bold red] {e}\n")
            return False

    def images_to_docx(self, image_files, output_path):
        """
        Convert images to DOCX using OCR

        Args:
            image_files: List of image file paths
            output_path: Output DOCX file path

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if not image_files:
                console.print("[red]No image files to convert![/red]")
                return False

            console.print(
                f"\n[bold cyan]Converting {len(image_files)} images to DOCX with OCR...[/bold cyan]")
            console.print(
                "[yellow]Note: This may take several minutes depending on image count.[/yellow]\n")

            doc = Document()
            sorted_files = sorted(image_files)

            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=console
            ) as progress:

                task = progress.add_task(
                    "[cyan]Processing images with OCR...",
                    total=len(sorted_files)
                )

                for idx, img_path in enumerate(sorted_files, 1):
                    try:
                        # Open and process image
                        img = Image.open(img_path)
                        text = pytesseract.image_to_string(
                            img,
                            lang=self.config.ocr_language
                        )

                        # Add to document
                        doc.add_heading(
                            f"Page {idx} - {os.path.basename(img_path)}", level=3)
                        for line in text.splitlines():
                            if line.strip():  # Only add non-empty lines
                                doc.add_paragraph(line)

                        progress.update(
                            task,
                            advance=1,
                            description=f"[green]Processed {os.path.basename(img_path)}"
                        )

                    except Exception as e:
                        console.print(
                            f"[yellow]Warning: Could not process {img_path}: {e}[/yellow]")
                        progress.update(task, advance=1)
                        continue

            # Save document
            doc.save(output_path)
            console.print(
                f"[bold green]✓ DOCX saved:[/bold green] {output_path}\n")
            return True

        except Exception as e:
            console.print(f"[bold red]✗ Error creating DOCX:[/bold red] {e}\n")
            return False

    def pdf_to_docx(self, pdf_path, output_path):
        """
        Convert PDF to DOCX (placeholder for future implementation)

        Args:
            pdf_path: Input PDF file path
            output_path: Output DOCX file path

        Returns:
            bool: True if successful, False otherwise
        """
        console.print(
            "[yellow]PDF to DOCX conversion requires additional libraries.[/yellow]")
        console.print(
            "[yellow]Consider using: pdf2docx or PyPDF2 + python-docx[/yellow]\n")
        return False
